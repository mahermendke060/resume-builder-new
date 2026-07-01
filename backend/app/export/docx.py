"""Render a tailored resume variant to an ATS-safe DOCX.

ATS-safe rules: single column, real heading paragraphs, no tables / text boxes /
headers-footers, standard fonts, plain bullet lists.
"""

import io
import re


def sanitize_text(text: str | None) -> str:
    if not text:
        return ""
    text = str(text)
    # Replace all dash variants
    text = re.sub(r"[\u2011\u2012\u2013\u2014\u2015]", "-", text)
    # Replace smart single quotes
    text = re.sub(r"[\u2018\u2019\u201A\u201B]", "'", text)
    # Replace smart double quotes
    text = re.sub(r"[\u201C\u201D\u201E\u201F]", '"', text)
    # Replace ellipsis
    text = re.sub(r"[\u2026]", "...", text)
    # Replace all space variants
    text = re.sub(r"[\u00A0\u2000-\u200A\u202F\u205F\u3000]", " ", text)
    # Remove control characters
    text = re.sub(r"[\u0000-\u001F\u007F-\u009F]", "", text)
    # Remove backticks
    text = text.replace("`", "")
    return text


def _bullet_text(b) -> str:
    return sanitize_text(b.get("text", "") if isinstance(b, dict) else str(b))


def render_docx(variant: dict, contact: dict | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    contact = contact or {}
    if contact.get("name"):
        h = doc.add_heading(sanitize_text(contact["name"]), level=0)
        h.alignment = 1  # center
    
    # Build contact line including links
    contact_parts = []
    if contact.get("email"):
        contact_parts.append(sanitize_text(contact["email"]))
    if contact.get("phone"):
        contact_parts.append(sanitize_text(contact["phone"]))
    if contact.get("location"):
        contact_parts.append(sanitize_text(contact["location"]))
    contact_line = " | ".join(contact_parts)
    if contact_line:
        p = doc.add_paragraph()
        p.add_run(contact_line)
    
    # Add links on new line
    if contact.get("links"):
        p = doc.add_paragraph()
        for i, link in enumerate(contact["links"]):
            if i > 0:
                p.add_run(" • ")
            p.add_run(sanitize_text(link))

    if variant.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(sanitize_text(variant["summary"]))

    if variant.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(sanitize_text(s) for s in variant["skills"]))

    if variant.get("experience"):
        doc.add_heading("Experience", level=1)
        for exp in variant["experience"]:
            title = " — ".join(x for x in (sanitize_text(exp.get("title")), sanitize_text(exp.get("company"))) if x)
            dates = " ".join(x for x in (sanitize_text(exp.get("start")), sanitize_text(exp.get("end"))) if x)
            p = doc.add_paragraph()
            p.add_run(title).bold = True
            if dates:
                p.add_run(f"  ({dates})")
            for b in exp.get("bullets", []) or []:
                doc.add_paragraph(_bullet_text(b), style="List Bullet")

    if variant.get("education"):
        doc.add_heading("Education", level=1)
        for ed in variant["education"]:
            line = ", ".join(
                sanitize_text(ed[k]) for k in ("degree", "institution", "year") if ed.get(k)
            )
            doc.add_paragraph(line)

    if variant.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for cert in variant["certifications"]:
            doc.add_paragraph(sanitize_text(cert), style="List Bullet")

    if variant.get("projects"):
        doc.add_heading("Projects", level=1)
        for proj in variant["projects"]:
            p = doc.add_paragraph()
            p.add_run(sanitize_text(proj.get("name", ""))).bold = True
            if proj.get("description"):
                doc.add_paragraph(sanitize_text(proj["description"]))
            if proj.get("bullets"):
                for b in proj["bullets"]:
                    doc.add_paragraph(_bullet_text(b), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
